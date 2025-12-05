#!/usr/bin/env python3
"""
PIF_Generator.py - Generates PIF sections based on country information from
Ass9 File Upload output files and Supabase database.
"""

import os
import sys
import json
import re
import requests
from pathlib import Path
from dataclasses import dataclass
from typing import Optional, Dict

# Try to import required libraries
try:
    import openai
except ImportError:
    print("Error: openai library not found. Please install it with: pip install openai")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx library not found. Please install it with: pip install python-docx")
    sys.exit(1)

try:
    from bs4 import BeautifulSoup
except ImportError:
    print("Warning: beautifulsoup4 library not found. HTML parsing for UNFCCC reports will be limited.")
    print("Install it with: pip install beautifulsoup4")
    BeautifulSoup = None

# Section specifications
@dataclass
class SectionSpec:
    key: str
    title: str
    word_limit: Optional[int] = None
    standard_text: Optional[str] = None
    prompt: Optional[str] = None
    keep_existing_prompt: bool = False

SECTIONS: Dict[str, SectionSpec] = {
    "rationale_intro": SectionSpec(
        key="rationale_intro",
        title="A. PROJECT RATIONALE",
        prompt=('''Reference if there is a past CBIT or GEF project in {Country} and its outcomes. If none, then leave blank.
                ''')
    ),
    "paris_etf": SectionSpec(
        key="paris_etf",
        title="The Paris Agreement and the Enhanced Transparency Framework",
        standard_text=(
            '''As part of the UNFCCC, the Paris Agreement (2015) strengthened the global response to climate change. 
            Article 13 established the Enhanced Transparency Framework (ETF), under which Parties report on mitigation, 
            adaptation and support. These information requirements present challenges to all countries, particularly those 
            already facing impacts.'''
        ),
        prompt=("Just keep standard text as is. Do not add additional paragraphs")
    ),
    "climate_transparency_country": SectionSpec(
        key="climate_transparency_country",
        title="Climate Transparency in {Country}",
        word_limit=350,
        prompt=(
            '''Explain where {Country} is not yet fully complying with ETF requirements, actions to date, and a 'without project' trajectory. 
                Identify drivers that sustain the status quo and motivate urgency. What are key drivers that would maintain the status quo (or make it worse)? E.g. population growth, economic development, climate change, socio-cultural and political factors, including conflicts, or technological changes. i.e. based on existing trends, the country would continue to not report in accordance with the ETF. Drivers are important as we have to respond to these in the description of the project
                Make sure to include specific numbers directly from these sections describing size, geographic features, inhabitants, and climate.
                - Create a paragraph describing its national framework, current reports (such as NIDs, NCDs, NCs, BURs, BTRs, and MRV system), include dates and other relevant figures. Make sure this updated to the latest and most current year.
                - Create another paragraph describing national strategy for energy transition and describe given information from the
                National Greenhouse Gas Inventory of its current emissions and current initiatives made under Paris Agreement to reduce GHG emissions. 
                Make sure to include specific statistics on emissions such as Carbon or Nitrogen, as well as what sectors emit the most GHGs)'''
        )
    ),
    "baseline_national_tf_header": SectionSpec(
        key="baseline_national_tf_header",
        title="1. National transparency framework",
        standard_text=(
            '''{Country} signed the UNFCCC on {UNFCCC_sign_date}, and ratified it on {UNFCCC_rat_date}. 
            It also ratified the Kyoto Protocol on {KP_rat_date}, and the Paris Agreement on {PA_rat_date}, 
            following its adoption in {PA_adopt_date}. The following sections describe {Country}'s institutional framework 
            for climate action, key legislation/policies, stakeholders, and ongoing transparency initiatives.'''
        ),
        prompt = ("Fill in the dates for UNFCCC, Kyoto Protocol, and Paris Agreement ratification/adoption and keep prompt as is. Do not change or add to standard text.")
    ),
    "baseline_institutional": SectionSpec(
        key="baseline_institutional",
        title="i. Institutional Framework for Climate Action",
        word_limit=500,
        prompt=('''
Structure and tone requirements:
• Maintain continuous paragraphs and professional language suitable for official reporting.
• **Use bullet points only for listing institutions and their roles — do NOT use Markdown tables or numbering.**
• The section should read as a formal, factual summary appropriate for inclusion in a GEF Project Identification Form.

Include the following content elements:

1. **Introductory paragraph**
   - Describe the national legal and institutional foundation for climate governance and transparency.
   - Mention the principal environmental or climate law that establishes the national climate framework, using the correct name and year where applicable.
   - Identify the lead ministry or agency responsible for coordinating national climate policy and reporting to the UNFCCC (e.g., the ministry of environment, science, or climate change).
   - Explain this institution's role in preparing National Communications, Biennial Update Reports (BURs), and Biennial Transparency Reports (BTRs).
    - Highlight role of President's office (if applicable) and any interministerial committee and describe any formalized process for engaging civil society.

2. **International support and collaboration**
   - Describe how international programs and initiatives—such as the Initiative for Climate Action Transparency (ICAT), the Capacity-Building Initiative for Transparency (CBIT), and the Global Environment Facility (GEF)—support institutional capacity building, MRV systems, and transparency processes.
   - Explain how these programs help align national MRV and reporting with the Enhanced Transparency Framework (ETF) of the Paris Agreement.

3. List the lead ministry, role of President's Office, interministerial committeee, and leading institutions and their roles using bullet points in this format:**
   Provide a comprehensive bullet list of institutions and their roles. Aim for multiple institutions (8–15+) with SHORT, CONCISE descriptions (typically 1 sentence per bullet).
   
   • Institution Name (Acronym)] [English translation]: [Brief 1-sentence description of primary role in climate action, MRV, or reporting].
   
   Examples (for Cuba — use appropriate institutions for the target country, with original language names and English translations in brackets):
   • Ministerio de Ciencia, Tecnología y Medio Ambiente (CITMA) [Ministry of Science, Technology, and Environment]: Lead ministry responsible for climate policy coordination and UNFCCC reporting.
   • Ministerio de Agricultura (MINAG) [Ministry of Agriculture]: Coordinates data collection and analysis for agriculture and LULUCF sectors.
   • Ministerio de Energía y Minas (MINEM) [Ministry of Energy and Mines]: Manages energy-related emissions monitoring and renewable energy transition.
   • Ministerio de Transporte (MITRANS) [Ministry of Transportation]: Oversees transport sector emissions data collection and mitigation strategies.
   • Ministerio de Industrias (MINDUS) [Ministry of Industries]: Coordinates industrial sector MRV and sustainable production practices.
   • Instituto Nacional de Recursos Hidráulicos (INRH) [National Institute of Hydraulic Resources]: Manages water resource management and climate adaptation data.
   • Ministerio de Finanzas y Precios [Ministry of Finance and Prices]: Oversees climate finance management and resource allocation for climate action.
   • Ministerio de Salud Pública (MINSAP) [Ministry of Public Health]: Contributes to health-climate nexus monitoring and adaptation planning.
   • Ministerio de Construcción (MICONS) [Ministry of Construction]: Coordinates building sector energy efficiency and climate resilience measures.
   • Ministerios de Educación y de Educación Superior (MINED/MES) [Ministries of Education]: Promote climate change awareness and environmental education.
   • Ministerio de Comercio Exterior e Inversión Extranjera (MINCEX) [Ministry of Foreign Trade and Foreign Investment]: Supports climate-related international cooperation and climate finance access.
Do not use the example if not applicable to {Country}.
The tone must be factual, cohesive, and formatted for direct PDF export with clear paragraph breaks and bullet list alignment.        
'''
                )
    ),
    "baseline_policy": SectionSpec(
        key="baseline_policy",
        title="ii. National Policy Framework",
        word_limit=500,
        prompt=(
            '''Write the section "National Policy Framework" for {Country} following the official GEF8 PIF format. This section focuses on specific LAWS, POLICIES, DECREES, and STRATEGIES—not on institutional frameworks or organizational structures.

Formatting and tone:
• Use an explicit bulleted list for the main policy and legal instruments. Each bullet must begin with a dash and a single space (e.g. - ), then the policy or law title followed by a colon, and then a short description of the instrument (maximum 3 sentences per bullet).
  - Example: - Decree 86 (2019): One-sentence summary of the decree's objectives. Second sentence on scope/focus areas if needed.
• CRITICAL: The opening paragraph and concluding paragraph must NEVER be in bullet point format. These are plain flowing paragraphs. DO NOT use bold formatting anywhere in this section.**
• Maintain a formal, factual tone suitable for inclusion in a GEF Project Identification Form (PIF).
• Information should be as specific as possible and drawn from the allowed sources.
• DO NOT include descriptions of institutional roles or organizational structures in this section; focus ONLY on the content, objectives, and scope of the laws, policies, and strategies themselves.
• After the bullet list, include a short concluding synthesis paragraph (1–2 paragraphs) summarizing the framework and outstanding challenges.

Content structure:

1. Opening paragraph (plain text paragraphs ONLY, NO BULLETS)
   - Write this as continuous flowing prose, NOT as bullet points.
   - DO NOT USE BULLETS HERE.
   - DO NOT use bold formatting anywhere.
   - Begin with one concise paragraph summarizing the country's overall policy and legal framework for climate action and how it aligns with international commitments such as the UNFCCC and the Paris Agreement.
   - End the paragraph with a transition like: "The following key instruments form the foundation of this framework:"

2. Policy and legal instruments (present as explicit bullets, EACH AS ITS OWN BULLET) — FOCUS ON LAWS AND POLICIES, NOT INSTITUTIONS
  - CRITICAL: Each law, decree, plan, or strategy MUST be output as a separate bullet line. DO NOT group multiple policies into one bullet.
  - For each law, decree, plan, or strategy, output a single bullet line in the following exact format:
    - [Full name and year]: [1–3 sentence description of objectives, focus areas (adaptation, mitigation, MRV, governance), and key provisions.]
  - Aim for 8–12+ policy/law instruments depending on the country. Make sure it is ordered chronologically and is updated to the current year of generation (>2025)
  - Examples (for illustration only — adapt names and years to the target country):
    - Law on Environmental Protection (Year): Establishes environmental protections and mandates for sustainable development. Addresses climate change impacts on key sectors.
    - National Climate Change Adaptation Plan (Year): Outlines priority sectors for climate adaptation including agriculture, water, and coastal zones. Sets targets and implementation mechanisms.
    - National Energy Transition Strategy (Year): Details targets for renewable energy deployment and energy efficiency improvements. Specifies timelines and sectoral focus areas.
    - National Determined Contribution (NDC) 3.0 (Year): Commits to specific greenhouse gas emission reduction targets by a set year. Identifies priority mitigation and adaptation actions.
    - Decree on Climate Governance (Year): Establishes institutional coordination mechanisms for climate policy implementation. Mandates integration of climate considerations in national planning.

3. Concluding synthesis paragraph (plain text paragraphs ONLY, NO BULLETS)
   - Write this as continuous flowing prose, NOT as bullet points.
   - Summarize:
     • The overall significance of this policy framework for national sustainability, climate resilience, and transparency.  
     • Persistent challenges such as limited technical capacity, financing gaps, or MRV system limitations.  
     • The need to regularly update policies and laws to ensure consistency with international commitments and the Enhanced Transparency Framework (ETF).
   - DO NOT use bold formatting anywhere.
   - DO NOT FORMAT this as bullet points; use flowing prose.
   '''
                )
    ),
    "baseline_stakeholders": SectionSpec(
        key="baseline_stakeholders",
        title="iii. Other key stakeholders for Climate Action",
        standard_text=("Non-government stakeholders for climate action are presented in Table 1."),
        prompt=(
        '''Extract non-government stakeholders and leverageable activities for {Country} from authoritative sources. PRIMARY: UNFCCC reports (NCs, BURs, BTRs) country pages; SECONDARY: PATPA, ICAT, NDC Partnership, GEF CBIT, official MRV/ministry pages.
        For EACH stakeholder category below, output entries as SEPARATE BULLETS using this exact format:
- Name: Existing activities

RULES:
- Max 8 entries per category.
- Each bullet = ONE entity only.
- "existing_activities" must be ≤200 chars, concise, factual, climate-relevant.
- Deduplicate entities across ALL categories.
- No marketing language. No filler.
- Activities must directly mention climate mitigation, adaptation, MRV, data, transparency, sector actions, or capacity-building.

===========================
REQUIRED CATEGORIES
===========================

1. Civil Society (CSOs and NGOs)
   - Name: Existing activities

2. Private sector
   - Name: Existing activities

3. Academia and research organizations
   - Name: Existing activities

4. Financial institutions / MDBs
   - Name: Existing activities

5. International organizations
   - Name: Existing activities

6. [Other – to be specified]
   - Name: Existing activities
each 'existing_activities' value ≤200 chars; deduplicate by name; concise and factual (no marketing language). Be pretty descriptive in the activities and its direct relevance.
   '''
        )
    ),
    "baseline_unfccc_reporting": SectionSpec(
        key="baseline_unfccc_reporting",
        title="iv. Official reporting to the UNFCCC",
        standard_text=("To meet its obligations under the UNFCCC, the country has submitted several documents related to its socio-economic development objectives (see Table 2)."),
        prompt = (
            '''Generate a bullet-point list of UNFCCC reports for {Country} using the scraped data provided below.

The UNFCCC reports have been scraped from unfccc.int/reports using the country identification number. 

YOUR TASK:
1. Extract all reports from the scraped data provided below
2. Format each report as a bullet point in the exact format: "- [Document Name]: [Date published]"
3. Sort all reports in chronological order by submission_date (earliest to latest)
4. Use the exact document names and dates as provided - do not modify them

OUTPUT FORMAT:
Return a bullet-point list where each bullet follows this exact format:

- [Document Name]: [Date published]
- [Document Name]: [Date published]
- [Document Name]: [Date published]

Example:
- Guinea-Bissau. National Communication (NC). NC 4.: 17 Nov 2025
- Guinea-Bissau. 2024 Biennial Transparency Report (BTR). BTR1. CTF-FTC: 14 Aug 2025
- Guinea-Bissau. 2024 National Inventory Document (NID): 26 Dec 2024

CRITICAL RULES:
- Use ONLY the scraped data provided below - do not invent or modify any information
- Keep document names exactly as scraped (no modifications to spelling, punctuation, or capitalization)
- Keep dates exactly as scraped (no format changes)
- Include ALL reports from the scraped data - do not skip any
- Sort by submission_date chronologically (earliest first)
- Format each entry as: "- [Document Name]: [Date published]" (with a dash, space, document name, colon, space, then date)
- If no scraped data is provided or scraping failed, return: "No UNFCCC reports found for this country."'''
        )
    ),
    "module_header": SectionSpec(
        key="module_header",
        title="2. Progress on the four Modules of the Enhanced Transparency Framework",
        standard_text=("The sections below outline status, progress, and challenges across the four core ETF modules."),
        prompt=("Just keep standard text as is. Do not add anything else.")
    ),
    "module_ghg": SectionSpec(
        key="module_ghg",
        title="i. GHG Inventory Module",
        prompt=(
        '''
        ** INITIAL PARAGRAPH **
        - Describe progress and gaps in GHG inventory (IPCC 2006, tiers, key categories, QA/QC, uncertainty, data systems, institutionalization).
        Summarize chronologically the country's GHG inventory submissions (National Communications, Biennial Update Reports, 
        National Inventory Reports) with years and data coverage.
        - Describe the institutional arrangement for inventory preparation—lead agency, technical team, and coordination with sectoral ministries and statistical offices.
        - Identify which IPCC Guidelines and methodological tiers are applied per sector (Energy, IPPU, Agriculture, LULUCF, Waste).
        - Highlight improvements achieved (adoption of 2006/2019 Guidelines, QA/QC systems, MRV platforms).
        - Present key challenges: data fragmentation, capacity constraints, absence of country-specific emission factors, staff turnover, and integration into planning.
        - Reference any ongoing capacity-building initiatives (e.g., CBIT, ICAT, PATPA projects).
        - Conclude with recommendations aligned with the Enhanced Transparency Framework (ETF):technical training (Tier 2/3, uncertainty analysis), 
        institutionalization of MRV protocols, linkage with NDC implementation and national plans.
        **NC/BTR/NID SUBMISSION SECTION**
        For {Country}, compile a list of its current Biennial Transparency Reports (BTR) / National Inventory Documents (NID or NIR), including the year uploaded and the document link (URL).

SOURCES TO SCRAPE (check all):

- https://unfccc.int/first-biennial-transparency-reports  
- https://unfccc.int/process-and-meetings/transparency-and-reporting/reporting-and-review-under-the-convention/national-communications-and-biennial-reports-annex-i-parties/biennial-report-submissions/second-biennial-reports-annex-i  
- https://unfccc.int/process-and-meetings/transparency-and-reporting/reporting-and-review-under-the-convention/national-communications-and-biennial-reports-annex-i-parties/biennial-report-submissions/third-biennial-reports-annex-i  
- https://unfccc.int/BR4  
- https://unfccc.int/BR5


You MUST extract entries from ALL of the following official pages:

1) First BTR page:
   https://unfccc.int/first-biennial-transparency-reports

2) BR2 submissions (Annex I):
   https://unfccc.int/process-and-meetings/transparency-and-reporting/reporting-and-review-under-the-convention/national-communications-and-biennial-reports-annex-i-parties/biennial-report-submissions/second-biennial-reports-annex-i

3) BR3 submissions (Annex I):
   https://unfccc.int/process-and-meetings/transparency-and-reporting/reporting-and-review-under-the-convention/national-communications-and-biennial-reports-annex-i-parties/biennial-report-submissions/third-biennial-reports-annex-i

4) BR4 submissions:
   https://unfccc.int/BR4

5) BR5 submissions:
   https://unfccc.int/BR5

6) For National Inventory Document (NID/NIR) links:
   You MUST return the direct link to the most recent NID/NIR available for {Country}
   (from any of the above pages or linked subpages).

INSTRUCTIONS:
STEPS:

1. Visit each URL above.  
2. Within each page, find rows or entries listing reports submitted by {Country}.  
    - Extract all BTR, BR2, BR3, BR4, BR5, NID/NIR entries relevant to {Country}.
3. For each matching BTR / NID / NIR entry, extract:  
   - Year of upload / submission (as YYYY), from the page info.  
4. If the report is a National Inventory Document (NID or NIR), clearly mark it as "NID" (or "NIR") in the output and link a direct link to the document.  
5. Ensure you capture all relevant entries across all the source pages.

- List results in reverse chronological order (newest first).
- If multiple documents exist in the same year, include all.
- If no information for a page exists, silently skip.

OUTPUT FORMAT (STRICT — NO VARIATION):
Return a clean bullet-point list in this exact form:

- <Report Type> (<YYYY>) — <URL>
- <Report Type> (<YYYY>) — <URL>
- ...
        '''
        )
    ),
    "module_adaptation": SectionSpec(
        key="module_adaptation",
        title="ii. Adaptation and Vulnerability Module",
        word_limit=400,
        prompt=('''State whether an Adaptation Communication has been submitted (and what year + link to source).
                - Summarize how {Country} prepares Adaption Communication, and explain how adaptation is integrated into national policy frameworks (e.g., National Adaptation Plan, state programs, NDC adaptation component, sectoral plans). Is there primarily a hired international consultant?
                - Highlight existing research and knowledge systems on vulnerability and risk (HVR studies, scenario development, academic programs).
                - Describe the status of monitoring and evaluation for adaptation actions and how it aligns with ETF paragraphs 104–117 of Decision 18/CMA.1.
                - describe what the country is doing to predict future climate changes and country vulnerabilities and detail the data systems in place to predict and where the data is from. Is there available information/data that could be drawn upon from national agencies?
                - Conclude with recommendations for strengthening adaptation governance, technical capacity, and integration with national planning systems. 
                '''
        )
    ),
    "module_ndc_tracking": SectionSpec(
        key="module_ndc_tracking",
        title="iii. NDC Tracking Module",
        word_limit=400,
        prompt=('''
TASK: Write a multi-paragraph analytical narrative for {Country} that closely matches the depth, structure, tone, and content style of the following EXAMPLE. The output should not copy the example's content, but must reproduce its format, level of specificity, and thematic organization.

===== STYLE + STRUCTURE TO FOLLOW (MUST MATCH EXAMPLE) =====

Generate ~6–7 paragraphs that cover the following components in order:

1. **NDC overview + scope of mitigation/adaptation actions**
   - Describe the breadth of actions (energy, transport, AFOLU, waste, water, industry, etc.).
   - Mention quantified or conditional/unconditional targets if they exist.
   - Note gaps in national systems for tracking NDC implementation.

2. **Summary of past CBIT / transparency / MRV strengthening efforts IF HAD PAST CBIT. Disregard section if none found or uploaded**
   - Identify key CBIT or transparency projects (CBIT 1, CBIT-AFOLU, CBIT-Global, or GEF-funded equivalents).
   - Describe tools, systems, templates, early NDC-tracking pilots, data hubs, or MRV components developed.
   - Explain which agency hosts these systems.

3. **Findings from the country's BTR1 (or BURs, NCs if BTR absent)**
   - Highlight persistent gaps: weak institutionalization, limited sectoral reporting, fragmented data flows.
   - State which ministries or institutions have insufficient mandates/capacity to track NDC actions.
   - Emphasize reliance on ad hoc reporting and manual consolidation.

4. **Systemic challenges identified (institutional fragmentation, sectoral silos, weak mandates)**
   - Describe coordination issues between lead transparency authority (e.g., EPA/Ministry of Environment) and cross-government bodies (e.g., planning commissions).
   - Note lack of enforcement authority, inconsistent indicator use, and gaps in national reporting frameworks.

5. **Subnational-level challenges (if relevant)**
   - Explain limited involvement of regional/local governments in NDC tracking.
   - Note absence of structured data pipelines from subnational actors.
   - Highlight particular weaknesses in adaptation monitoring at local levels.

6. **Forward-looking needs and recommendations**
   - Identify urgent needs to strengthen institutional arrangements, indicator harmonization, data flows, reporting mandates, and MRV system integration with national planning.
   - Link these needs to ETF requirements and BTR obligations.

===== CONTENT REQUIREMENTS =====
- Write in a formal, analytical, policy-oriented tone.
- Use sector names, institutions, and terminology relevant to transparency, MRV, NDC tracking, and planning.
- Do NOT invent detailed numerical targets unless typical for {Country}; instead, generalize with correct structure.
- When referring to institutions, use placeholders if unknown (e.g., "the national environment authority", "the planning commission", "sector ministries").
- The narrative must be country-specific but inference-based; no fictional project names unless plausible analogs (e.g., "CBIT project on MRV enhancement", "National Climate Data System").
- NO bullet points — continuous prose only.

===== OUTPUT FORMAT =====
Return ONLY the final narrative text with paragraphs separated by line breaks. No introductions, no explanations, no notes, no markdown, no metadata.
Overall, the output should highlight if the country is tracking implementation of NDC. How it tracks its progress, if it does not have any systemmatic way of tracking progress, etc.
''')
    ),    
    "module_support": SectionSpec(
        key="module_support",
        title="iv. Support Needed and Received Module",
        word_limit=400,
        prompt=(
            '''
            TASK: {Country} should provide information on national circumstances and institutional arrangements relevant to reporting on support needed and received, including: (a) A description of the systems and processes used to identify, track and report support needed and received, including a description of the challenges and limitations; (b) Information on country priorities and strategies and on any aspects of the Party's NDC under Article 4 of the Paris Agreement that need support.
            ASK: Write a multi-paragraph analytical narrative for {Country} that mirrors the tone, structure, and depth of the EXAMPLE text below. 
DO NOT copy the example. Instead, reproduce its style, sequencing, and level of institutional and diagnostic detail.

===== CONTENT + STRUCTURE TO FOLLOW (MUST MATCH EXAMPLE) =====

Write 5–7 paragraphs covering the following in order:

1. **Importance of support tracking + national circumstances**
   - Explain why tracking support needed/received is essential for transparency, planning, national priorities, and accessing climate finance.
   - Summarize national circumstances relevant to support tracking (institutional capacity, climate commitments, NDC structure).
   - If available, mention high-level cost estimates or financing needs for NDC implementation (conditional/unconditional).

2. **Current systems, institutional arrangements, and ETF-relevant requirements**
   - Describe existing systems/processes used to identify, track, and report support needed and received.
   - Include: institutional arrangements, coordinating bodies, MRV authorities, ministries involved (e.g., MoF, Planning Commission, Environment Authority).
   - Address ETF requirement (a): "A description of the systems and processes used to identify, track, and report support needed and received, including challenges and limitations."

3. **Results and limitations of past CBIT, MRV, or transparency projects**
   - Summarize what CBIT or equivalent transparency projects achieved (templates, early systems, data hubs, pilot tools).
   - Note what did NOT get institutionalized (lack of scaling, fragmented adoption, weak mandates, no nationwide reporting system).
   - Describe remaining capacity gaps per terminal evaluations or BTR findings.

4. **Findings from BTR1 (or NID, BURs if BTR absent) on support reporting**
   - Highlight what support data is captured (multilateral/bilateral funds, technical assistance, capacity-building).
   - Emphasize gaps: disaggregation issues, lack of methodology, no aggregation system, incomplete alignment with NDC priorities.
   - Explain challenges in verifying data, tracking off-budget support, and engaging all stakeholders.

5. **National challenges: institutional fragmentation, weak mandates, separate systems**
   - Describe how ministries operate in silos, lack harmonized indicators, and maintain separate reporting cycles and formats.
   - Highlight gaps in coordination among finance, planning, and environment authorities.
   - Explain the absence of mechanisms for regular, structured reporting by sectors or subnational actors.

6. **Country priorities + NDC support needs (ETF requirement b)**
   - Summarize the country's priority sectors needing support and areas of the NDC requiring financial, technological, or capacity-building assistance.
   - Describe barriers such as limited knowledge of funding opportunities, insufficient tracking of expenditures, or lack of standardized tools.
   - Explain why a national climate finance MRV system is needed to meet ETF obligations.

7. **Forward-looking conclusion**
   - Identify critical needs for institutionalizing support tracking, harmonizing systems, improving mandates, and integrating MRV with national planning.
   - Summarize why robust support MRV is essential for ETF reporting, resource mobilization, and long-term planning.

===== STYLE REQUIREMENTS =====
- Formal, analytical, policy-focused tone.
- Multi-paragraph prose (NO bullets, NO lists).
- Frequent reference to institutions, systems, MRV, climate finance, NDC implementation, and transparency frameworks.
- Use sector ministries, planning bodies, and environment authorities generically unless real names are known.
- DO NOT invent highly specific numerical targets unless the country is known to publicly report them; otherwise generalize.
'''
            ),
    ),
    "other_baseline_initiatives": SectionSpec(
        key="other_baseline_initiatives",
        title="Other baseline initiatives",
        standard_text=("This CBIT project is aligned with and complements other initiatives supported by the GEF and development partners in the Country, as outlined in Table below"),
        prompt = (''' 
        Fill the table with relevant transparency initiatives for {Country}.
        TASK: For {Country}, fill out the following table with REAL transparency- and MRV-related initiatives currently ongoing or recently completed, using ONLY authoritative sources (UNFCCC, GEF/CBIT, ICAT, NDC Partnership, PMI, REDD+, SDGs, national MRV portals, and official government documentation).

The blue items in the first column below are EXAMPLES of types of programs/projects. 
Replace them with the ACTUAL programs/projects relevant to {Country}. 
If a category does not exist for {Country}, skip it silently.

===========================
TABLE TO POPULATE (REQUIRED)
===========================

For each relevant initiative in {Country}, fill all fields:

1. Program / Project  
2. Leading ministry and supporting entities  
3. Brief description (≤180 characters, factual, no marketing)  
4. Duration (start and end year)  
5. Value (USD) — if unknown, write "N/A"  
6. Relation to ETF and transparency system (describe EXACT contribution to MRV, NDC tracking, inventory, reporting, etc.)

Use the following rows (these are EXAMPLES of categories, not content):
- First Biennial Transparency Report and latest National Communication  
- CBIT 1 Project  
- ICAT  
- NDC Partnership  
- Sustainable Development Goals (SDGs)  
- Partnership for Market Implementation (PMI)  
- REDD+  
- Any additional transparency/MRV-relevant national projects

===========================
OUTPUT FORMAT (STRICT)
===========================

Return a strict RFC-8259 JSON object with key "body" whose value is a STRING containing:

{
  "table_data": [
    {
      "program_project": "<string>",
      "leading_entities": "<string>",
      "description": "<string>",
      "duration": "<string>",
      "value_usd": "<string>",
      "relation_to_ETF": "<string>"
    }
  ]
}

Rules:
- Include 5–12 initiatives depending on availability.
- ALL fields must be filled; no empty strings.
- Use precise years from project documentation.
- Deduplicate by program name.
- Use only verified information; no speculative entries.
        ''' 
        )
    ),
    "key_barriers": SectionSpec(
        key="key_barriers",
        title="Key barriers",
        prompt=(
            '''Summarize barriers around organizing climate data (Component 1), incomplete ETF modules/capacity (Component 2), 
            and reliance on projects/external consultants with limited use in planning (Components 1 & 3).'''
            )
    ),
    "barrier1": SectionSpec(
        key="Barrier1",
        word_limit=200,
        title="Barrier 1: {Country} lacks the capacity to systematically organize climate data",
        prompt=(
            '''This barrier corresponds to the project's component 1. (1-2 paragraphs)
            Describe {Country}'s lack of a national (or strong) climate transparency system and insufficient institutional arrangements, 
            procedures, and protocols to allow for the collection of required data.'''
                )
    ),
    "barrier2": SectionSpec(
        key="barrier2",
        word_limit=200,
        title="Barrier 2: {Country}'s climate ETF modules for GHG Inventory, adaptation/vulnerability, NDC tracking, and support needed and received are incomplete and not fully aligned with ETF requirements.",
        prompt=(
            '''This barrier corresponds to the project's component 2. (1-2 paragraphs)
            Describe {Country}'s limited technical content and related capacity on the four ETF chapters
            (mitigation/inventory, adaptation/vulnerability modelling, NDC tracking (both mitigation and adaptation),
             tracking of support needed and received'''
                )
    ),
    "barrier3": SectionSpec(
        key="barrier3",
        word_limit=200,
        title="Barrier 3: {Country} lacks capacity to consistently use its climate change information for reporting to the UNFCCC and for national planning without project-based financing and external consultants.",
        standard_text=('''
        While {Country} has demonstrated strong commitment by submitting multiple NCs, BURs, [and most recently its first BTR1], 
        the country's reporting system relies heavily on project-based support and external expertise. 
        National reporting processes are not yet institutionalized or adequately funded through government budgets, 
        making them vulnerable to discontinuity once donor-funded projects conclude. This creates a dependency loop in which reporting quality 
        and frequency are linked to the availability of external financing and consultants, rather than sustained national capacity.
        Furthermore, climate data generated for international reporting is not routinely integrated into national planning or 
        development decision-making. Ministries and planning agencies often lack the tools, capacity, or incentives to use transparency 
        outputs—such as emissions data or climate finance tracking—in their sectoral strategies or policy formulation. 
        This weakens the feedback loop between transparency and implementation, reducing the impact of climate information on real-world outcomes.
        The lack of institutional ownership and mainstreamed use of transparency findings across the policy landscape undermines 
        {Country}'s ability to align its national development priorities with its climate commitments. 
        Strengthening internal capacity, operationalizing the national transparency platform for dual reporting and planning functions,
         and embedding transparency workflows into regular government systems are therefore essential for long-term sustainability and 
        compliance with ETF obligations.''' 
        )
    ),
}

# Supabase configuration
SUPABASE_URL = "https://tulbxwdifnzquliytsog.supabase.co"
SUPABASE_API_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InR1bGJ4d2RpZm56cXVsaXl0c29nIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NjI5OTg0MTAsImV4cCI6MjA3ODU3NDQxMH0.pRnak9Ii7Eqli-o8AEYX0DCyaWOi04OlEhLoynw88wU"

def validate_openai_api_key(api_key):
    """
    Validate an OpenAI API key by making a test API call.
    Returns True if valid, False otherwise.
    """
    if not api_key or not api_key.strip():
        return False
    
    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key.strip())
        
        # Make a simple test call to validate the key
        response = client.models.list()
        return True
    except Exception as e:
        # If there's any error, the key is invalid
        return False

def get_openai_api_key():
    """
    Prompt user for OpenAI API key with validation.
    Returns the API key string if valid.
    """
    while True:
        user_input = input("Please provide an OpenAI API key: ").strip()
        
        if not user_input:
            print("API key cannot be empty. Please try again.")
            continue
        
        # Validate the API key
        print("Validating API key...")
        if validate_openai_api_key(user_input):
            print("✓ Valid API key.")
            return user_input
        else:
            print("✗ Invalid API key. Please try again.")

def get_cookie_information():
    """
    Prompt user for UNFCCC cookie information and save to JSON file.
    Returns True if cookies were provided and saved, False otherwise.
    """
    print("\n" + "="*80)
    print("UNFCCC Cookie Information (Optional)")
    print("="*80)
    print("To scrape UNFCCC reports, you may need to provide session cookies.")
    print("If you skip this step, scraping will proceed without authentication.")
    print("\nTo get your cookies:")
    print("1. Open your browser and navigate to https://unfccc.int/reports")
    print("2. Open Developer Tools (F12)")
    print("3. Go to Application/Storage tab > Cookies > https://unfccc.int")
    print("4. Copy the cookie names and values")
    print("\nPress Enter to skip, or type 'yes' to provide cookies:")
    
    response = input().strip().lower()
    
    if response not in ['yes', 'y']:
        print("Skipping cookie setup. Scraping will proceed without authentication.")
        return False
    
    cookies = {}
    headers = {}
    
    print("\nEnter cookie information (press Enter with empty name to finish):")
    cookie_count = 0
    while True:
        cookie_name = input(f"Cookie name #{cookie_count + 1} (or press Enter to finish): ").strip()
        if not cookie_name:
            break
        
        cookie_value = input(f"Cookie value for '{cookie_name}': ").strip()
        if cookie_value:
            cookies[cookie_name] = cookie_value
            cookie_count += 1
        else:
            print("Cookie value cannot be empty. Skipping this cookie.")
    
    # Ask about custom headers
    print("\nDo you have any custom headers to add? (e.g., Authorization, X-CSRF-Token)")
    header_response = input("Type 'yes' to add headers, or press Enter to skip: ").strip().lower()
    
    if header_response in ['yes', 'y']:
        header_count = 0
        while True:
            header_name = input(f"Header name #{header_count + 1} (or press Enter to finish): ").strip()
            if not header_name:
                break
            
            header_value = input(f"Header value for '{header_name}': ").strip()
            if header_value:
                headers[header_name] = header_value
                header_count += 1
            else:
                print("Header value cannot be empty. Skipping this header.")
    
    # Save to JSON file
    if cookies or headers:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        json_path = os.path.join(script_dir, "unfccc_cookies.json")
        
        cookie_data = {
            "cookies": cookies,
            "headers": headers
        }
        
        try:
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(cookie_data, f, indent=2, ensure_ascii=False)
            print(f"\n✓ Cookie information saved to {json_path}")
            print(f"  - {len(cookies)} cookie(s) saved")
            if headers:
                print(f"  - {len(headers)} header(s) saved")
            return True
        except Exception as e:
            print(f"\n✗ Error saving cookie information: {e}")
            return False
    else:
        print("\nNo cookie information provided.")
        return False

def search_output_files(country_name, output_folder):
    """
    Search for country-related files in the Ass9 File Upload Output folder.
    Returns list of file paths and their contents.
    """
    country_name_lower = country_name.lower()
    found_files = []
    
    if not os.path.exists(output_folder):
        print(f"Warning: Output folder {output_folder} does not exist.")
        return found_files
    
    # Search for files containing country name
    for file_path in Path(output_folder).glob('*'):
        if file_path.is_file():
            filename_lower = file_path.name.lower()
            if country_name_lower in filename_lower:
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    found_files.append({
                        'filename': file_path.name,
                        'content': content
                    })
                    print(f"Found file: {file_path.name}")
                except Exception as e:
                    print(f"Error reading {file_path.name}: {e}")
    
    return found_files

def get_country_data_from_supabase(country_name):
    """
    Query Supabase database for country information.
    Returns list of matching country records with their sections data.
    """
    try:
        # Query Supabase REST API
        url = f"{SUPABASE_URL}/rest/v1/countries"
        headers = {
            "apikey": SUPABASE_API_KEY,
            "Authorization": f"Bearer {SUPABASE_API_KEY}",
            "Content-Type": "application/json",
            "Prefer": "return=representation"
        }
        
        # Try to get all countries first, then filter
        # This handles cases where 'names' might be JSON or text
        try:
            # First try with ilike filter (if names is a text column)
            params = {
                "select": "*",
                "names": f"ilike.%{country_name}%"
            }
            response = requests.get(url, headers=headers, params=params, timeout=30)
            response.raise_for_status()
            countries = response.json()
        except:
            # If that fails, get all records and filter manually
            response = requests.get(url, headers=headers, params={"select": "*"}, timeout=30)
            response.raise_for_status()
            countries = response.json()
        
        # Filter by country name (case-insensitive)
        matching_countries = []
        country_name_lower = country_name.lower()
        
        for country in countries:
            # Check if country name matches (case-insensitive)
            country_names = country.get('names', '')
            matched = False
            
            if isinstance(country_names, str):
                if country_name_lower in country_names.lower():
                    matched = True
            elif isinstance(country_names, list):
                for name in country_names:
                    if country_name_lower in str(name).lower():
                        matched = True
                        break
            
            # Also check if country name appears anywhere in the record
            if not matched:
                country_str = json.dumps(country).lower()
                if country_name_lower in country_str:
                    matched = True
            
            if matched and country not in matching_countries:
                matching_countries.append(country)
        
        return matching_countries
    
    except Exception as e:
        print(f"Error querying Supabase: {e}")
        import traceback
        traceback.print_exc()
        return []

def load_cookies_from_json(json_path=None):
    """
    Load cookies and session data from a JSON file.
    
    Expected JSON format:
    {
        "cookies": {
            "cookie_name_1": "cookie_value_1",
            "cookie_name_2": "cookie_value_2"
        },
        "headers": {
            "header_name_1": "header_value_1"
        }
    }
    
    Returns: tuple of (cookies_dict, headers_dict)
    """
    if json_path is None:
        # Try to find cookies.json in the same directory as the script
        script_dir = os.path.dirname(os.path.abspath(__file__))
        json_path = os.path.join(script_dir, "unfccc_cookies.json")
    
    cookies = {}
    custom_headers = {}
    
    if os.path.exists(json_path):
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                cookies = data.get("cookies", {})
                custom_headers = data.get("headers", {})
                print(f"Loaded cookies and headers from {json_path}")
        except Exception as e:
            print(f"Warning: Could not load cookies from {json_path}: {e}")
            print("Continuing without cookies...")
    else:
        print(f"Note: Cookie file not found at {json_path}")
        print("Scraping will proceed without authentication cookies.")
        print("If the site requires authentication, create unfccc_cookies.json with your session cookies.")
    
    return cookies, custom_headers

def get_country_reports_by_id(country_id):
    """
    Scrape UNFCCC reports using country ID (based on scrape_unfccc.py).
    Uses the corporate_author filter with country_id.
    
    Returns a list of dictionaries with keys: 'name', 'submission_date'
    """
    try:
        # Load cookies and custom headers from JSON file
        cookies, custom_headers = load_cookies_from_json()
        
        base_url = "https://unfccc.int/reports"
        
        # Headers to mimic a browser request
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Accept-Encoding": "gzip, deflate, br",
            "Connection": "keep-alive",
            "Upgrade-Insecure-Requests": "1",
            "Referer": "https://unfccc.int/"
        }
        
        # Merge custom headers from JSON
        headers.update(custom_headers)
        
        # Filter by corporate author using country_id
        params = {
            "f[0]": f"corporate_author:{country_id}",
            "items_per_page": 50,   # big enough to avoid "Load more" for now
            "view": "table",        # list view with the table
        }
        
        response = requests.get(base_url, params=params, headers=headers, cookies=cookies, timeout=30)
        response.raise_for_status()
        
        if BeautifulSoup:
            soup = BeautifulSoup(response.text, "html.parser")
            
            # Find the main results table
            table = soup.find("table")
            if not table:
                print(f"    Warning: Could not find results table on the page")
                return []
            
            tbody = table.find("tbody")
            if not tbody:
                # Try to find rows directly in table
                rows = table.find_all("tr")[1:]  # Skip header
            else:
                rows = tbody.find_all("tr")
            
            results = []
            
            for tr in rows:
                tds = tr.find_all("td")
                if len(tds) < 4:
                    # Unexpected format, skip
                    continue
                
                # UNFCCC table layout (at time of writing) is:
                # [0] Document name, [1] Type of document, [2] Author, [3] Submission date
                name = tds[0].get_text(strip=True)
                submission_date = tds[3].get_text(strip=True) if len(tds) > 3 else ""
                
                if name:  # Only add if we have a document name
                    results.append({
                        "name": name,
                        "submission_date": submission_date
                    })
            
            return results
        else:
            print("    Warning: BeautifulSoup not available, cannot parse HTML")
            return []
    
    except Exception as e:
        print(f"    Error scraping UNFCCC reports with country_id {country_id}: {e}")
        import traceback
        traceback.print_exc()
        return []

def scrape_unfccc_reports(country_name):
    """
    Scrape UNFCCC reports page (unfccc.int/reports) for a specific country.
    Extracts document name, type, and submission date from the results table.
    Also returns the raw HTML for AI processing.
    
    Uses cookies from unfccc_cookies.json if available for authentication.
    
    Returns a tuple: (list of dictionaries with keys: 'document_name', 'type', 'submission_date', raw_html_string)
    """
    unfccc_reports = []
    raw_html = ""
    
    try:
        # Load cookies and custom headers from JSON file
        cookies, custom_headers = load_cookies_from_json()
        
        # URL for UNFCCC reports search
        base_url = "https://unfccc.int/reports"
        
        # Headers to mimic a browser request
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Accept-Encoding": "gzip, deflate, br",
            "Connection": "keep-alive",
            "Upgrade-Insecure-Requests": "1",
            "Referer": "https://unfccc.int/"
        }
        
        # Merge custom headers from JSON (these will override defaults)
        headers.update(custom_headers)
        
        # Try different URL patterns and methods
        search_attempts = [
            # GET requests with query parameters - prioritize author search
            {"method": "GET", "url": f"{base_url}?author={country_name}"},
            {"method": "GET", "url": f"{base_url}?search={country_name}"},
            {"method": "GET", "url": f"{base_url}?country={country_name}"},
            {"method": "GET", "url": f"{base_url}?q={country_name}"},
            # Try the base URL and search in HTML
            {"method": "GET", "url": base_url},
        ]
        
        for attempt in search_attempts:
            try:
                if attempt["method"] == "GET":
                    # Use cookies if available
                    response = requests.get(attempt["url"], headers=headers, cookies=cookies, timeout=30)
                else:
                    continue
                
                response.raise_for_status()
                
                # Store raw HTML for AI processing (use the first successful response with results)
                if not raw_html or len(unfccc_reports) == 0:
                    raw_html = response.text
                
                # Parse HTML if BeautifulSoup is available
                if BeautifulSoup:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    
                    # Look for table with reports - try multiple selectors
                    tables = soup.find_all('table')
                    
                    # Also try to find divs with table-like structure
                    if not tables:
                        # Some sites use div-based tables
                        table_divs = soup.find_all('div', class_=re.compile(r'table|results|report', re.I))
                        for div in table_divs:
                            # Check if it contains rows
                            rows = div.find_all(['tr', 'div'], class_=re.compile(r'row|item|result', re.I))
                            if rows:
                                tables.append(div)
                    
                    for table in tables:
                        # Find all rows (skip header row if it's a table element)
                        if table.name == 'table':
                            rows = table.find_all('tr')[1:]  # Skip header
                        else:
                            # For div-based tables, find all row-like elements
                            rows = table.find_all(['tr', 'div'], class_=re.compile(r'row|item|result', re.I))
                        
                        for row in rows:
                            # Try to find cells (td, th, or div with cell-like classes)
                            cells = row.find_all(['td', 'th'])
                            if not cells:
                                # Try div-based cells
                                cells = row.find_all('div', class_=re.compile(r'cell|col|field', re.I))
                            
                            if len(cells) >= 3:  # At least document name, type, date
                                try:
                                    # Extract text from cells
                                    cell_texts = [cell.get_text(strip=True) for cell in cells]
                                    
                                    # Try to identify which cell is which based on content
                                    # Usually: document_name, type, author, submission_date
                                    document_name = cell_texts[0] if len(cell_texts) > 0 else ""
                                    doc_type = cell_texts[1] if len(cell_texts) > 1 else ""
                                    author = cell_texts[2] if len(cell_texts) > 2 else ""
                                    submission_date = cell_texts[3] if len(cell_texts) > 3 else ""
                                    
                                    # If we have 4+ cells, assume: name, type, author, date
                                    # If we have 3 cells, assume: name, type, date (no author column)
                                    if len(cell_texts) == 3:
                                        document_name = cell_texts[0]
                                        doc_type = cell_texts[1]
                                        submission_date = cell_texts[2]
                                        author = country_name  # Assume it matches if we're searching for it
                                    
                                    # Include ALL documents where author matches country
                                    # This is the primary filter - if author matches, include the document
                                    doc_name_lower = document_name.lower()
                                    author_lower = author.lower() if author else ""
                                    country_lower = country_name.lower()
                                    
                                    # Primary filter: author must match country (case-insensitive, partial match)
                                    # This ensures we capture ALL documents for the country
                                    author_matches = author_lower and country_lower in author_lower
                                    
                                    # Secondary filter: if no author field or author doesn't match, 
                                    # check if country name appears in document name
                                    country_in_name = country_lower in doc_name_lower
                                    
                                    # Include if author matches OR (no author field and country in name) OR (3 cells - likely no author column)
                                    # This ensures we get ALL significant documents for the country
                                    should_include = author_matches or (not author and country_in_name) or len(cell_texts) == 3
                                    
                                    if should_include:
                                        # Only add if we have meaningful data (document name is required)
                                        if document_name and document_name.strip():
                                            unfccc_reports.append({
                                                'document_name': document_name,
                                                'type': doc_type,
                                                'submission_date': submission_date
                                            })
                                except (IndexError, AttributeError) as e:
                                    continue
                    
                    # Continue trying other URLs to get all possible results
                    # Don't break early - combine results from all attempts
                else:
                    # Fallback: try to extract from raw HTML using regex
                    # Look for table-like patterns
                    table_pattern = r'<tr[^>]*>.*?<td[^>]*>(.*?)</td>.*?<td[^>]*>(.*?)</td>.*?<td[^>]*>(.*?)</td>(?:.*?<td[^>]*>(.*?)</td>)?'
                    matches = re.findall(table_pattern, response.text, re.DOTALL | re.IGNORECASE)
                    
                    for match in matches:
                        if len(match) >= 3:
                            doc_name = re.sub(r'<[^>]+>', '', match[0]).strip()
                            doc_type = re.sub(r'<[^>]+>', '', match[1]).strip()
                            if len(match) >= 4:
                                author = re.sub(r'<[^>]+>', '', match[2]).strip()
                                date = re.sub(r'<[^>]+>', '', match[3]).strip()
                            else:
                                author = ""
                                date = re.sub(r'<[^>]+>', '', match[2]).strip()
                            
                            # Include ALL documents where author matches country
                            doc_name_lower = doc_name.lower()
                            author_lower = author.lower() if author else ""
                            country_lower = country_name.lower()
                            
                            # Primary filter: author must match country (case-insensitive, partial match)
                            # This ensures we capture ALL documents for the country
                            author_matches = author_lower and country_lower in author_lower
                            
                            # Secondary filter: if no author field or author doesn't match,
                            # check if country name appears in document name
                            country_in_name = country_lower in doc_name_lower
                            
                            # Include if author matches OR (no author and country in name)
                            # This ensures we get ALL significant documents for the country
                            if author_matches or (not author and country_in_name):
                                if doc_name and doc_name.strip():
                                    unfccc_reports.append({
                                        'document_name': doc_name,
                                        'type': doc_type,
                                        'submission_date': date
                                    })
            
            except requests.RequestException:
                # Try next URL
                continue
            except Exception as e:
                # Continue to next URL
                continue
        
        # Remove duplicates based on document_name
        seen = set()
        unique_reports = []
        for report in unfccc_reports:
            doc_name = report.get('document_name', '').strip()
            if doc_name and doc_name not in seen:
                seen.add(doc_name)
                unique_reports.append(report)
        
        # Sort by submission date (most recent first) if dates are available
        def sort_key(report):
            date_str = report.get('submission_date', '')
            # Try to extract year for sorting
            year_match = re.search(r'(\d{4})', date_str)
            if year_match:
                return -int(year_match.group(1))  # Negative for descending
            return 0
        
        unique_reports.sort(key=sort_key)
        
        # Debug: print summary of what we found
        if unique_reports:
            print(f"    Scraped {len(unique_reports)} unique report(s)")
            # Count by type
            type_counts = {}
            for report in unique_reports:
                doc_type = report.get('type', 'Unknown')
                type_counts[doc_type] = type_counts.get(doc_type, 0) + 1
            print(f"    Report types found: {dict(type_counts)}")
        
        return unique_reports, raw_html
    
    except Exception as e:
        print(f"Error scraping UNFCCC reports: {e}")
        import traceback
        traceback.print_exc()
        return [], ""

def extract_sections_from_country_data(country_data_list):
    """
    Extract all relevant sections from the country data sections.
    Returns a dictionary mapping section keys to their extracted data.
    Handles the Supabase format: {"sections": [{"name": "...", "documents": [...]}]}
    """
    # Mapping of section keys to possible name patterns found in Supabase
    # Patterns can be strings or lists, will be normalized to lists
    section_name_patterns = {
        'rationale_intro': ['A. PROJECT RATIONALE', 'PROJECT RATIONALE', 'rationale', 'cbit project', 'gef project'],
        'paris_etf': ['The Paris Agreement and the Enhanced Transparency Framework', 'Paris Agreement', 'Enhanced Transparency Framework', 'ETF'],
        'climate_transparency_country': ['Climate Transparency', 'climate transparency', 'transparency in'],
        'baseline_national_tf_header': ['National transparency framework', 'National Transparency Framework', '1. National transparency framework'],
        'baseline_institutional': ['Institutional Framework for Climate Action', 'Institutional framework for climate action', 'i. Institutional Framework', 'Institutional Framework'],
        'baseline_policy': ['National Policy Framework', 'National policy framework', 'ii. National Policy Framework', 'Policy Framework'],
        'baseline_stakeholders': ['Other key stakeholders for Climate Action', 'Other key stakeholders', 'iii. Other key stakeholders', 'key stakeholders', 'stakeholders'],
        'baseline_unfccc_reporting': ['Official Reporting to the UNFCCC', 'Official reporting to the UNFCCC', 'iv. Official reporting', 'UNFCCC reporting'],
        'module_ghg': ['GHG Inventory', 'GHG inventory', 'i. GHG Inventory', 'GHG Inventory Module', 'greenhouse gas inventory'],
        'module_adaptation': ['Adaptation and Vulnerability', 'Adaptation and vulnerability', 'ii. Adaptation and Vulnerability', 'Adaptation Module'],
        'module_ndc_tracking': ['NDC Tracking', 'NDC tracking', 'iii. NDC Tracking', 'NDC Tracking Module'],
        'module_support': ['Support Needed and Received', 'Support needed and received', 'iv. Support Needed and Received', 'Support Module'],
        'other_baseline_initiatives': ['Other Baseline Initiatives', 'Other baseline initiatives', 'Baseline Initiatives', 'baseline initiatives'],
        'key_barriers': ['Key Barriers', 'Key barriers', 'barriers'],
        'barrier1': ['Barrier 1', 'barrier 1'],
        'barrier2': ['Barrier 2', 'barrier 2'],
        'barrier3': ['Barrier 3', 'barrier 3'],
    }
    
    # Initialize sections_data dictionary with all section keys
    sections_data = {key: [] for key in section_name_patterns.keys()}
    
    for country_data in country_data_list:
        sections = country_data.get('sections', [])
        
        # Handle JSON string format
        if isinstance(sections, str):
            try:
                sections = json.loads(sections)
            except:
                sections = []
        
        # Handle case where sections might be nested in a dict
        if isinstance(sections, dict) and 'sections' in sections:
            sections = sections['sections']
        
        if not isinstance(sections, list):
            continue
        
        for section in sections:
            section_name = section.get('name', '')
            if not section_name:
                continue
            
            section_name_lower = section_name.lower()
            matched = False
            
            # Try to match section to one of our patterns
            for section_key, patterns in section_name_patterns.items():
                # Normalize patterns to always be a list
                if isinstance(patterns, str):
                    patterns = [patterns]
                
                for pattern in patterns:
                    pattern_lower = pattern.lower()
                    # Try exact match first, then substring match
                    if pattern_lower == section_name_lower or pattern_lower in section_name_lower or section_name_lower in pattern_lower:
                        sections_data[section_key].append(section)
                        matched = True
                        break
                
                if matched:
                    break
            
            # Fallback: Check section key or id if available
            if not matched:
                section_key_field = section.get('key', '').lower() or section.get('id', '').lower()
                for section_key in section_name_patterns.keys():
                    if section_key.lower() in section_key_field or section_key_field in section_key.lower():
                        sections_data[section_key].append(section)
                        break
    
    return sections_data

def format_sections_text(sections_data):
    """
    Format the extracted sections data into readable text.
    Handles all sections including those that may contain tables.
    """
    formatted_text = ""
    
    # Map section keys to display names
    section_display_names = {
        'rationale_intro': 'A. PROJECT RATIONALE',
        'paris_etf': 'The Paris Agreement and the Enhanced Transparency Framework',
        'climate_transparency_country': 'Climate Transparency',
        'baseline_national_tf_header': 'National transparency framework',
        'baseline_institutional': 'Institutional framework for climate action',
        'baseline_policy': 'National policy framework',
        'baseline_stakeholders': 'Other key stakeholders for Climate Action',
        'baseline_unfccc_reporting': 'Official Reporting to the UNFCCC',
        'module_ghg': 'GHG Inventory',
        'module_adaptation': 'Adaptation and Vulnerability',
        'module_ndc_tracking': 'NDC Tracking',
        'module_support': 'Support Needed and Received',
        'other_baseline_initiatives': 'Other Baseline Initiatives',
    }
    
    for section_key, section_list in sections_data.items():
        if section_list:
            # Use display name if available, otherwise use key
            display_name = section_display_names.get(section_key, section_key.replace('_', ' ').title())
            # Handle case where display_name might be a list - use first item
            if isinstance(display_name, list):
                display_name = display_name[0]
            formatted_text += f"\n=== {display_name} ===\n"
            
            for section in section_list:
                documents = section.get('documents', [])
                if documents:
                    for doc in documents:
                        doc_type = doc.get('doc_type', 'Unknown')
                        extracted_text = doc.get('extracted_text', '')
                        
                        # Check if this looks like table data (JSON, structured data)
                        if extracted_text:
                            # Try to detect if it's JSON table data
                            try:
                                table_data = json.loads(extracted_text)
                                if isinstance(table_data, dict) and 'table_data' in table_data:
                                    # Format as table
                                    formatted_text += f"\n[From {doc_type} - Table Data]:\n"
                                    formatted_text += format_table_data(table_data['table_data'])
                                    formatted_text += "\n"
                                elif isinstance(table_data, list) and len(table_data) > 0 and isinstance(table_data[0], dict):
                                    # List of dictionaries - format as table
                                    formatted_text += f"\n[From {doc_type} - Table Data]:\n"
                                    formatted_text += format_table_data(table_data)
                                    formatted_text += "\n"
                                else:
                                    # Regular text content
                                    formatted_text += f"\n[From {doc_type}]:\n{extracted_text}\n"
                            except (json.JSONDecodeError, TypeError):
                                # Not JSON, treat as regular text
                                formatted_text += f"\n[From {doc_type}]:\n{extracted_text}\n"
                else:
                    # If no documents, try to get text directly from section
                    section_text = section.get('text', '') or section.get('content', '')
                    
                    # Check if section has table_data field
                    table_data = section.get('table_data', None)
                    if table_data:
                        formatted_text += "\n[Table Data]:\n"
                        if isinstance(table_data, str):
                            try:
                                table_data = json.loads(table_data)
                            except:
                                pass
                        if isinstance(table_data, (list, dict)):
                            formatted_text += format_table_data(table_data)
                            formatted_text += "\n"
                    
                    if section_text:
                        formatted_text += f"\n{section_text}\n"
            formatted_text += "\n"
    
    return formatted_text

def format_table_data(table_data):
    """
    Format table data (list of dictionaries) into a readable table format.
    """
    if not table_data:
        return ""
    
    # Handle different table data structures
    if isinstance(table_data, dict):
        # If it's a dict with 'table_data' key, extract it
        if 'table_data' in table_data:
            table_data = table_data['table_data']
        else:
            # Convert single dict to list
            table_data = [table_data]
    
    if not isinstance(table_data, list) or len(table_data) == 0:
        return ""
    
    # Get all unique keys from all rows
    all_keys = set()
    for row in table_data:
        if isinstance(row, dict):
            all_keys.update(row.keys())
    
    if not all_keys:
        return ""
    
    # Convert to sorted list for consistent ordering
    headers = sorted(all_keys)
    
    # Build table
    table_lines = []
    
    # Header row
    header_row = " | ".join(str(h).title() for h in headers)
    table_lines.append(header_row)
    table_lines.append("-" * len(header_row))
    
    # Data rows
    for row in table_data:
        if isinstance(row, dict):
            values = [str(row.get(key, ""))[:50] for key in headers]  # Limit cell length
            table_lines.append(" | ".join(values))
    
    return "\n".join(table_lines)

def read_section_examples():
    """
    Read the Section Examples.txt file.
    """
    script_dir = os.path.dirname(os.path.abspath(__file__))
    examples_path = os.path.join(script_dir, 'Section Examples.txt')
    
    try:
        with open(examples_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Warning: Could not read Section Examples.txt: {e}")
        return ""

def format_section_title(title, country_name):
    """Format section title by replacing {Country} placeholder."""
    return title.replace("{Country}", country_name)

def format_standard_text(standard_text, country_name, **kwargs):
    """Format standard text by replacing placeholders."""
    formatted = standard_text.replace("{Country}", country_name)
    for key, value in kwargs.items():
        placeholder = "{" + key + "}"
        if placeholder in formatted:
            formatted = formatted.replace(placeholder, str(value) if value else "[Not available]")
    return formatted

def should_keep_standard_text_only(section_key):
    """Check if this section should use standard text only without additional content.
    Only 'paris_etf' and 'module_header' sections use standard_text only."""
    return section_key in ["paris_etf", "module_header"]

def generate_single_section(api_key, country_name, section_spec, output_files_content, supabase_sections_text, section_examples):
    """
    Generate a single section based on its specification.
    Returns the generated text or standard text if applicable.
    """
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    
    # Handle standard text sections
    if section_spec.standard_text:
        # Only use standard text alone for specific sections (paris_etf and module_header)
        # All other sections should generate additional content even if they have standard_text
        if should_keep_standard_text_only(section_spec.key):
            # Return standard text with placeholders filled
            return format_standard_text(section_spec.standard_text, country_name)
        
        # For all other sections with standard_text, we'll generate additional content
        # The standard_text will be included as part of the generation instructions
    
    # If no standard text or customization needed, generate from scratch
    # Prepare section-specific prompt
    section_prompt = section_spec.prompt or ""
    
    # Replace {Country} placeholder in prompt
    section_prompt = section_prompt.replace("{Country}", country_name)
    
    # Limit content to fit within token limits
    total_char_limit = 100000  # Conservative limit per section
    base_prompt_length = 5000
    available_chars = total_char_limit - base_prompt_length
    
    # Truncate content if needed
    output_files_len = len(output_files_content)
    supabase_len = len(supabase_sections_text)
    examples_len = len(section_examples)
    total_content_len = output_files_len + supabase_len + examples_len
    
    truncated_output = output_files_content
    truncated_supabase = supabase_sections_text
    truncated_examples = section_examples
    
    if total_content_len > available_chars:
        scale_factor = available_chars / total_content_len
        if output_files_len > 0:
            truncated_output = output_files_content[:int(output_files_len * scale_factor * 0.4)] + "\n[... truncated ...]"
        if supabase_len > 0:
            truncated_supabase = supabase_sections_text[:int(supabase_len * scale_factor * 0.3)] + "\n[... truncated ...]"
        if examples_len > 0:
            truncated_examples = section_examples[:int(examples_len * scale_factor * 0.3)] + "\n[... truncated ...]"
    
    # Build the full prompt
    word_limit_text = f"MINIMUM length: {section_spec.word_limit} words, GOAL: approximately {section_spec.word_limit} words." if section_spec.word_limit else ""
    
    # Scrape UNFCCC reports data for baseline_unfccc_reporting section using country_id
    unfccc_reports_data = ""
    if section_spec.key == "baseline_unfccc_reporting":
        print(f"    Scraping UNFCCC reports for {country_name}...")
        print(f"    Please provide the UNFCCC country identification number for {country_name}")
        print(f"    (e.g., 442 for Guinea-Bissau - you can find this by inspecting the UNFCCC reports page)")
        country_id = input(f"    Enter UNFCCC country ID for {country_name} (or press Enter to skip): ").strip()
        
        if country_id:
            print(f"    Scraping with country_id: {country_id}...")
            reports = get_country_reports_by_id(country_id)
            if reports:
                print(f"    Found {len(reports)} UNFCCC report(s)")
                # Format the scraped data for the prompt
                unfccc_reports_data = "\n\n=== UNFCCC Reports Data (scraped from unfccc.int/reports) ===\n"
                unfccc_reports_data += f"The following reports were scraped for {country_name} (country_id: {country_id}):\n\n"
                for i, report in enumerate(reports, 1):
                    unfccc_reports_data += f"{i}. Document Name: {report.get('name', 'N/A')}\n"
                    unfccc_reports_data += f"   Submission Date: {report.get('submission_date', 'N/A')}\n\n"
            else:
                unfccc_reports_data = "\n\n=== UNFCCC Reports Data (from unfccc.int/reports) ===\n"
                unfccc_reports_data += f"No reports found for country_id {country_id}.\n\n"
        else:
            print(f"    Skipping UNFCCC scraping (no country_id provided)")
            unfccc_reports_data = "\n\n=== UNFCCC Reports Data (from unfccc.int/reports) ===\n"
            unfccc_reports_data += "No country ID provided. Cannot scrape UNFCCC reports.\n\n"
    
    standard_text_instruction = ""
    if section_spec.standard_text:
        if section_spec.key == "baseline_national_tf_header":
            standard_text_instruction = f"\n\nSTANDARD TEXT TO USE (fill in dates):\n{format_standard_text(section_spec.standard_text, country_name)}\n\nUse this standard text structure and fill in the missing dates ({{UNFCCC_sign_date}}, {{UNFCCC_rat_date}}, {{KP_rat_date}}, {{PA_rat_date}}, {{PA_adopt_date}}) based on the information provided. Replace the placeholders with actual dates found in the provided information."
        elif section_spec.key == "baseline_stakeholders":
            standard_text_instruction = f"\n\nSTANDARD TEXT TO INCLUDE AT THE END:\n{section_spec.standard_text}\n\nAfter generating the stakeholder content, append this standard text."
        elif section_spec.key == "baseline_unfccc_reporting":
            standard_text_instruction = f"\n\nSTANDARD TEXT TO INCLUDE AT THE END:\n{section_spec.standard_text}\n\nAfter generating the UNFCCC reporting table, append this standard text."
        elif section_spec.key == "other_baseline_initiatives":
            standard_text_instruction = f"\n\nSTANDARD TEXT TO INCLUDE AT THE BEGINNING:\n{section_spec.standard_text}\n\nStart with this standard text, then generate the table content."
    
    all_info = f"""
COUNTRY: {country_name}

=== Information from Ass9 File Upload Output Files ===
{truncated_output}

=== Information from Supabase Database ===
{truncated_supabase}{unfccc_reports_data}

=== EXAMPLE SECTIONS (Reference Only - These are example answers showing desired format, style, and level of detail) ===
The following examples demonstrate how the sections should be written. Use these as a reference for:
- Paragraph structure and flow
- Level of detail and explanation
- Professional tone and factual presentation
- How to integrate quantitative and qualitative information
- Format for tables and structured content

{truncated_examples}
"""
    
    full_prompt = f"""You are an expert at drafting PIF (Project Identification Form) sections for climate transparency projects.

CRITICAL: Focus on FACTUALITY and ACCURACY. Reduce creativity. Base everything strictly on the provided information.

The "=== EXAMPLE SECTIONS (Reference Only) ===" section above contains EXAMPLE ANSWERS that demonstrate the desired format, style, paragraph structure, and level of detail. Use these examples as your primary reference for how to structure and write your sections.

{word_limit_text}

SECTION TO GENERATE: {section_spec.title}

INSTRUCTIONS FOR THIS SECTION:
{section_prompt}
{standard_text_instruction}

STRICT REQUIREMENTS:
- Use the information provided above as the PRIMARY source. You may supplement with information from PATPA and ICAT sources or UNEP trusted sources if needed, but ALL such references MUST be properly cited.
- Do not invent, infer, or speculate beyond what is explicitly stated in sources.
- If information is missing, clearly state what is missing or what gaps exist - do not fill gaps with assumptions.
- Follow the format, style, and structure requirements specified in the instructions above.
- Include ALL relevant quantitative data (numbers, amounts, dates, percentages) exactly as they appear in the sources.
- Include ALL relevant qualitative information (descriptions, assessments, challenges, achievements).
- When referencing projects/initiatives from ICAT and PATPA documents, cite them properly.
- Write in a professional, factual tone - avoid creative language or embellishment.
- Prioritize accuracy and completeness.

Generate this section now, ensuring maximum detail and factual accuracy:"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert at drafting PIF sections for climate transparency projects. Your primary focus is FACTUALITY and ACCURACY. You extract and synthesize information from provided sources without adding creative elements. You strictly adhere to the information provided and do not invent or speculate."},
                {"role": "user", "content": full_prompt}
            ],
            max_tokens=4000,
            temperature=0.1
        )
        
        return response.choices[0].message.content
    
    except Exception as e:
        print(f"Error calling OpenAI API for section {section_spec.key}: {e}")
        return None

def generate_all_sections(api_key, country_name, output_files_content, supabase_sections_text, section_examples):
    """
    Generate all sections defined in SECTIONS dictionary.
    Returns a dictionary mapping section keys to generated content.
    """
    generated_sections = {}
    
    # Define the order of sections
    section_order = [
        "rationale_intro",
        "paris_etf",
        "climate_transparency_country",
        "baseline_national_tf_header",
        "baseline_institutional",
        "baseline_policy",
        "baseline_stakeholders",
        "baseline_unfccc_reporting",
        "module_header",
        "module_ghg",
        "module_adaptation",
        "module_ndc_tracking",
        "module_support",
        "other_baseline_initiatives",
        "key_barriers",
        "barrier1",
        "barrier2",
        "barrier3",
    ]
    
    print(f"\nGenerating {len(section_order)} sections...")
    
    for i, section_key in enumerate(section_order, 1):
        if section_key not in SECTIONS:
            print(f"Warning: Section key '{section_key}' not found in SECTIONS dictionary. Skipping.")
            continue
        
        section_spec = SECTIONS[section_key]
        print(f"  [{i}/{len(section_order)}] Generating: {section_spec.title}")
        
        # Handle standard text sections that should only use standard text
        # Only paris_etf and module_header use standard_text alone
        if section_spec.standard_text and should_keep_standard_text_only(section_key):
            generated_text = format_standard_text(section_spec.standard_text, country_name)
            generated_sections[section_key] = generated_text
            print(f"    ✓ Using standard text only")
            continue
        
        # Generate the section
        generated_text = generate_single_section(
            api_key,
            country_name,
            section_spec,
            output_files_content,
            supabase_sections_text,
            section_examples
        )
        
        if generated_text:
            generated_sections[section_key] = generated_text
            print(f"    ✓ Generated successfully")
        else:
            print(f"    ✗ Failed to generate")
            generated_sections[section_key] = f"[ERROR: Failed to generate {section_spec.title}]"
    
    return generated_sections

def main():
    # Step 1: Get country name
    country_name = input("What country do you want to draft the PIF for: ").strip()
    
    if not country_name:
        print("Error: Country name cannot be empty.")
        return
    
    print(f"\nProcessing PIF generation for {country_name}...")
    
    # Step 2: Search for files in Ass9 File Upload Output folder
    script_dir = os.path.dirname(os.path.abspath(__file__))
    ass9_output_folder = os.path.join(script_dir, '..', 'Ass9 File Upload', 'Output')
    ass9_output_folder = os.path.abspath(ass9_output_folder)
    
    print(f"\nSearching for files in: {ass9_output_folder}")
    output_files = search_output_files(country_name, ass9_output_folder)
    
    # Combine output file contents
    output_files_content = ""
    if output_files:
        for file_data in output_files:
            output_files_content += f"\n--- From {file_data['filename']} ---\n{file_data['content']}\n"
    else:
        output_files_content = "[No matching files found in Ass9 File Upload Output folder]"
        print("No matching files found in Ass9 File Upload Output folder.")
    
    # Step 3: Query Supabase for country data
    print(f"\nQuerying Supabase database for {country_name}...")
    country_data_list = get_country_data_from_supabase(country_name)
    
    if country_data_list:
        print(f"Found {len(country_data_list)} matching country record(s) in Supabase.")
    else:
        print("No matching country records found in Supabase database.")
    
    # Step 4: Extract sections from Supabase data
    print(f"\nExtracting sections from Supabase data...")
    sections_data = extract_sections_from_country_data(country_data_list)
    
    # Print summary of extracted sections
    for section_key, section_list in sections_data.items():
        if section_list:
            display_name = section_list[0].get('name', section_key) if section_list else section_key
            print(f"  Found {len(section_list)} entry(ies) for: {display_name}")
    
    supabase_sections_text = format_sections_text(sections_data)
    
    if not supabase_sections_text.strip():
        supabase_sections_text = "[No section data found in Supabase database]"
        print("  No section data extracted from Supabase.")
    else:
        print(f"  Formatted {sum(len(v) for v in sections_data.values())} section entries from Supabase.")
    
    # Step 5: Read section examples
    section_examples = read_section_examples()
    
    # Step 6: Get OpenAI API key
    print("\n" + "="*80)
    api_key = get_openai_api_key()
    
    # Step 7: Get UNFCCC cookie information (optional)
    get_cookie_information()
    
    # Step 8: Generate all sections with AI
    print("\nGenerating PIF sections with AI...")
    generated_sections = generate_all_sections(
        api_key,
        country_name,
        output_files_content,
        supabase_sections_text,
        section_examples
    )
    
    if not generated_sections:
        print("Error: Failed to generate sections.")
        return
    
    # Step 9: Format and save all sections to Word document
    output_folder = os.path.join(script_dir, 'Output')
    os.makedirs(output_folder, exist_ok=True)
    
    output_filename = f"{country_name} section draft.docx"
    output_path = os.path.join(output_folder, output_filename)
    
    # Define section order for output
    section_order = [
        "rationale_intro",
        "paris_etf",
        "climate_transparency_country",
        "baseline_national_tf_header",
        "baseline_institutional",
        "baseline_policy",
        "baseline_stakeholders",
        "baseline_unfccc_reporting",
        "module_header",
        "module_ghg",
        "module_adaptation",
        "module_ndc_tracking",
        "module_support",
        "other_baseline_initiatives",
        "key_barriers",
        "barrier1",
        "barrier2",
        "barrier3",
    ]
    
    # Create Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading(f"PIF SECTIONS FOR {country_name.upper()}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add each section in order
    sections_written = 0
    for section_key in section_order:
        if section_key not in generated_sections:
            continue
        
        section_spec = SECTIONS.get(section_key)
        if not section_spec:
            continue
        
        # Format section title
        section_title = format_section_title(section_spec.title, country_name)
        section_content = generated_sections[section_key]
        
        # Remove section title from content if it appears at the beginning
        # This prevents headers from being duplicated in the body
        section_content_lines = section_content.split('\n')
        if section_content_lines:
            first_line = section_content_lines[0].strip()
            # Remove markdown headers
            first_line_clean = re.sub(r'^#{1,6}\s+', '', first_line)
            # Check if first line matches section title (case-insensitive, ignoring formatting)
            title_normalized = re.sub(r'[^\w\s]', '', section_title.lower()).strip()
            first_line_normalized = re.sub(r'[^\w\s]', '', first_line_clean.lower()).strip()
            if first_line_normalized == title_normalized or first_line_normalized.startswith(title_normalized[:20]):
                # Remove the first line if it matches the title
                section_content = '\n'.join(section_content_lines[1:]).strip()
        
        # Add section heading
        heading = doc.add_heading(section_title, level=1)
        
        # Add section content with formatting
        add_formatted_content(doc, section_content)
        
        # Add spacing after section
        doc.add_paragraph()
        sections_written += 1
    
    # Save document
    doc.save(output_path)
    
    print(f"\n✓ Output file created: {output_path}")
    print(f"  Generated {sections_written} sections for: {country_name}")

def clean_bullet_text(text):
    """
    Clean bullet text to remove duplicate patterns like "**Category**Category" or full content duplication.
    Handles cases like entire bullet content being repeated.
    Examples:
    - "Content.Content" -> "Content"
    - "**Policy (2018)**: Description.**Policy (2018)**: Description." -> "**Policy (2018)**: Description."
    """
    import re
    
    if not text:
        return text
    
    # Quick check: if text length is even and substantial, check if first half == second half
    text_len = len(text)
    if text_len > 50 and text_len % 2 == 0:
        mid = text_len // 2
        first_half = text[:mid].strip()
        second_half = text[mid:].strip()
        
        # Normalize and compare
        first_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', first_half.lower()))
        second_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', second_half.lower()))
        
        if len(first_norm) > 30 and first_norm == second_norm:
            return first_half.strip()
    
    # First, detect and remove full content duplication
    # The pattern is usually: "Content.Content" or "Content Content" where entire content repeats
    
    # Method 1: Check if text appears to be duplicated by splitting at likely break points
    text_len = len(text)
    if text_len > 40:  # Only check if text is substantial
        # Look for common break points (periods, colons followed by space)
        break_points = []
        for i, char in enumerate(text):
            if char in ['.', ':'] and i < text_len - 1 and text[i+1] in [' ', '\n']:
                break_points.append(i + 1)
        
        # Check each break point to see if text after it duplicates text before it
        for break_point in break_points:
            if break_point < text_len * 0.3 or break_point > text_len * 0.7:
                continue  # Skip if break is too early or too late
            
            first_part = text[:break_point].strip()
            second_part = text[break_point:].strip()
            
            if len(first_part) < 20 or len(second_part) < 20:
                continue
            
            # Normalize for comparison (remove punctuation, extra spaces, case)
            first_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', first_part.lower())).strip()
            second_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', second_part.lower())).strip()
            
            # Check if second part is very similar to first part
            if len(first_norm) > 30 and len(second_norm) > 30:
                # Check if first 40+ chars match
                match_length = min(40, len(first_norm), len(second_norm))
                if first_norm[:match_length] == second_norm[:match_length]:
                    # Very likely duplicate - return first part
                    return first_part.strip()
                
                # Also check if second part contains most of first part
                if len(first_norm) > 50:
                    first_chunk = first_norm[:min(50, len(first_norm))]
                    if first_chunk in second_norm[:min(100, len(second_norm))]:
                        # Likely duplicate
                        return first_part.strip()
    
    # Method 2: Look for repeated sequences separated by period (with or without space)
    # Pattern: "Long text.Long text" where entire phrase repeats
    # This handles cases like "Content.Content" (no space after period)
    
    # Try to find where text repeats by looking for the start of text appearing again
    if text_len > 40:
        # Find potential start points of duplication (look for start of text appearing again)
        first_30_chars = text[:30].strip()
        first_50_chars = text[:50].strip()
        
        # Look for these patterns appearing again in the text
        for search_text in [first_50_chars, first_30_chars]:
            if len(search_text) < 20:
                continue
            
            # Normalize search text
            search_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', search_text.lower())).strip()
            
            # Find where this pattern appears again (should be around middle/end)
            for i in range(len(text) // 2, len(text) - 10):
                candidate = text[i:i+len(search_text)+20].strip()
                candidate_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', candidate.lower())).strip()
                
                if len(candidate_norm) > 20 and search_norm[:min(30, len(search_norm))] in candidate_norm[:min(50, len(candidate_norm))]:
                    # Found potential duplicate - check if text before i is similar to text after i
                    before_text = text[:i].strip()
                    after_text = text[i:].strip()
                    
                    before_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', before_text.lower())).strip()
                    after_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', after_text.lower())).strip()
                    
                    if len(before_norm) > 30 and len(after_norm) > 30:
                        # Compare first 40 chars
                        compare_len = min(40, len(before_norm), len(after_norm))
                        if before_norm[:compare_len] == after_norm[:compare_len]:
                            # Duplicate found - return first part
                            return before_text.strip()
    
    # Method 3: Look for pattern "Text.Text" where Text repeats
    # Try different sequence lengths
    for seq_len in range(min(200, text_len // 2), 40, -10):
        # Pattern: long sequence + period (optional space) + same sequence
        pattern = r'(.{' + str(seq_len) + r',}?)(\.\s*|\s+)(\1)'
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if match:
            first_seq = match.group(1).strip()
            second_seq = match.group(3).strip()
            
            # Normalize and compare
            first_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', first_seq.lower())).strip()
            second_norm = re.sub(r'\s+', ' ', re.sub(r'[^\w\s]', '', second_seq.lower())).strip()
            
            if len(first_norm) > 30:
                # Check if they match (first 40 chars should match)
                match_chars = min(40, len(first_norm), len(second_norm))
                if first_norm[:match_chars] == second_norm[:match_chars]:
                    # Remove the duplicate
                    text = text.replace(match.group(0), first_seq, 1)
                    break
    
    # Remove pattern: "**Text**: somethingText: something" -> "**Text**: something"
    # Handles cases like "**First National Communication (NC)**: 201First National Communication (NC): 201"
    text = re.sub(r'\*\*([^*]+)\*\*\s*:?\s*([^:]*?)\1(\s|$|:|\.)', r'**\1**: \2', text, flags=re.IGNORECASE)
    
    # Remove pattern: "**Text**Text" or "**Text** Text" -> "**Text**"
    # Handles cases like "**Private sector**Private sector"
    text = re.sub(r'\*\*([^*]+)\*\*\s*\1(\s|$|:|\.)', r'**\1**\2', text, flags=re.IGNORECASE)
    text = re.sub(r'\*\*([^*]+)\*\*\s*\*\*\1\*\*', r'**\1**', text, flags=re.IGNORECASE)
    
    # Remove pattern with numbers: "2. **Category**Category" -> "Category"
    text = re.sub(r'\d+\.\s+\*?\*?([^*\d]+)\*?\*?\s*\1', r'\1', text, flags=re.IGNORECASE)
    
    # Remove pattern: "**Name:**Name:" or "**Name:**Name" -> "Name:"
    text = re.sub(r'\*\*([^*:]+):\*\*\s*\1:?\s*', r'\1: ', text, flags=re.IGNORECASE)
    
    # Remove duplicate phrases (case-insensitive, handles multi-word phrases)
    # Pattern: "Private sectorPrivate sector" -> "Private sector"
    text = re.sub(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\s+\1\b', r'\1', text, flags=re.IGNORECASE)
    
    # Removed duplicate word regex as it was causing spelling issues
    
    return text.strip()

def clean_content(content):
    """
    Clean content by removing markdown headers, duplicate titles, and formatting issues.
    """
    import re
    
    # Remove markdown headers (##, ###, etc.) from start of lines
    content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
    
    lines = content.split('\n')
    cleaned_lines = []
    
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Skip completely empty lines (they'll be handled separately)
        if not line_stripped:
            cleaned_lines.append(line)
            continue
        
        # Skip markdown headers
        if re.match(r'^#{1,6}\s+', line_stripped):
            continue
        
        # Remove duplicate patterns like "**Private sector**Private sector" or "2. **Private sector**Private sector"
        # Pattern 1: Remove duplicate after bold formatting (handles "**Text**Text" or "**Text**: somethingText: something")
        # This handles cases like "**First National Communication (NC)**: 201First National Communication (NC): 201"
        line_stripped = re.sub(r'\*\*([^*]+)\*\*\s*:?\s*([^:]*?)\1(\s|$|:)', r'**\1**: \2', line_stripped, flags=re.IGNORECASE)
        # Pattern 1b: Simple case "**Text**Text" -> "**Text**"
        line_stripped = re.sub(r'\*\*([^*]+)\*\*\s*\1(\s|$|:|\.)', r'**\1**\2', line_stripped, flags=re.IGNORECASE)
        # Pattern 2: Remove duplicate when both have bold
        line_stripped = re.sub(r'\*\*([^*]+)\*\*\s*\*\*\1\*\*', r'**\1**', line_stripped, flags=re.IGNORECASE)
        # Pattern 3: Handle numbered duplicates like "2. **Category**Category"
        line_stripped = re.sub(r'(\d+\.\s+)?\*?\*?([^*]+)\*?\*?\s*\2', r'\2', line_stripped, flags=re.IGNORECASE)
        
        # Remove standalone category headers that are just numbers + category name
        # Check if this line is like "2. **Category**" or "2. Category" followed by bullets
        if re.match(r'^\d+\.\s+(\*?\*?[^*]+\*?\*?)\s*$', line_stripped):
            # Check if next non-empty line is a bullet - if so, skip this category header
            skip_category_header = False
            for j in range(i + 1, len(lines)):
                next_line = lines[j].strip()
                if not next_line:
                    continue
                if next_line.startswith('- ') or next_line.startswith('• ') or next_line.startswith('* '):
                    # This is a category header before bullets, skip it
                    skip_category_header = True
                    break
                else:
                    # Next line is not a bullet, so this might be content we want to keep
                    break
            
            if skip_category_header:
                continue
        
        # For bullet points, keep them as-is (no regex reformatting)
        if line_stripped.startswith('- ') or line_stripped.startswith('• ') or line_stripped.startswith('* '):
            # No reformatting - keep bullet text exactly as it appears
            cleaned_lines.append(line)
            continue
        
        cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def add_formatted_content(doc, content):
    """
    Add formatted content to a Word document.
    Handles paragraphs, bullet points, tables, and basic formatting.
    """
    import re
    
    if not content:
        return
    
    # Clean content first - remove markdown headers and duplicates
    content = clean_content(content)
    
    # First check if content contains JSON table data
    json_data = extract_json_table_data_with_metadata(content)
    if json_data:
        table_data = json_data.get('table_data')
        summary = json_data.get('summary')
        
        if table_data:
            # Create Word table from JSON data
            create_word_table(doc, table_data)
        
        # Add summary if present
        if summary:
            doc.add_paragraph()
            summary_para = doc.add_paragraph(summary)
        
        # Remove the JSON from content and process remaining text
        content = remove_json_from_content(content)
        if not content.strip():
            return
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    current_paragraph = None
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Skip markdown headers
        if re.match(r'^#{1,6}\s+', line):
            i += 1
            continue
        
        # Check for bullet points (lines starting with - or •)
        if line.startswith('- ') or line.startswith('• ') or line.startswith('* '):
            # Start a bullet list
            bullet_text = line[2:].strip() if len(line) > 2 else line[1:].strip()
            
            # No regex reformatting - use bullet text as-is
            paragraph = doc.add_paragraph(bullet_text, style='List Bullet')
            add_inline_formatting(paragraph, bullet_text)
            current_paragraph = None
            
            # Add consecutive bullet points
            i += 1
            while i < len(lines) and lines[i].strip() and (lines[i].strip().startswith('- ') or 
                                     lines[i].strip().startswith('• ') or 
                                     lines[i].strip().startswith('* ')):
                bullet_text = lines[i].strip()[2:].strip() if len(lines[i].strip()) > 2 else lines[i].strip()[1:].strip()
                # No regex reformatting - use bullet text as-is
                paragraph = doc.add_paragraph(bullet_text, style='List Bullet')
                add_inline_formatting(paragraph, bullet_text)
                i += 1
        
        # Check for numbered lists (lines starting with number.)
        # But skip if it looks like a category header before bullets
        elif re.match(r'^\d+\.\s+', line):
            # Check if this is a category header (short, followed by bullets)
            next_non_empty = None
            for j in range(i + 1, min(i + 3, len(lines))):
                if lines[j].strip():
                    next_non_empty = lines[j].strip()
                    break
            
            # If next line is a bullet, this is likely a category header - skip it
            if next_non_empty and (next_non_empty.startswith('- ') or 
                                  next_non_empty.startswith('• ') or 
                                  next_non_empty.startswith('* ')):
                i += 1
                continue
            
            # Otherwise, treat as numbered list item
            list_text = re.sub(r'^\d+\.\s+', '', line)
            # No regex reformatting - use list text as-is
            paragraph = doc.add_paragraph(list_text, style='List Number')
            add_inline_formatting(paragraph, list_text)
            current_paragraph = None
            i += 1
        
        # Regular paragraph - accumulate multiple lines until empty line
        else:
            # Clean duplicate text patterns before adding to paragraph
            # Handle "**Text**: somethingText: something" -> "**Text**: something"
            line = re.sub(r'\*\*([^*]+)\*\*\s*:?\s*([^:]*?)\1(\s|$|:|\.)', r'**\1**: \2', line, flags=re.IGNORECASE)
            # Handle "**Text**Text" -> "**Text**"
            line = re.sub(r'\*\*([^*]+)\*\*\s*\1(\s|$|:|\.)', r'**\1**\2', line, flags=re.IGNORECASE)
            # Removed duplicate word regex as it was causing spelling issues
            
            # Skip if line looks like a header that's already been added as a section title
            # Check if line is just a title (short, no punctuation, might be all caps or title case)
            if len(line) < 100 and not re.search(r'[.!?]', line) and (line.isupper() or line.istitle()):
                # Check if this looks like a section header that shouldn't be in body
                # Skip standalone headers that are likely duplicates
                if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s*$', line):
                    # This might be a header, but we'll keep it if it's part of content
                    # Only skip if it's very short and looks like a title
                    if len(line.split()) <= 5:
                        # Check next line - if it's empty or starts a new section, this might be a duplicate header
                        if i + 1 < len(lines):
                            next_line = lines[i + 1].strip()
                            if not next_line or next_line.startswith('#') or next_line.startswith('- ') or next_line.startswith('• '):
                                i += 1
                                continue
            
            if current_paragraph is None:
                current_paragraph = doc.add_paragraph()
                add_inline_formatting(current_paragraph, line)
            else:
                # Add space and continue the paragraph
                current_paragraph.add_run(' ')
                add_inline_formatting(current_paragraph, line)
            i += 1
    
    # If there's a trailing paragraph, ensure it's closed
    if current_paragraph is not None:
        pass  # Already added

def extract_json_table_data_with_metadata(content):
    """
    Extract JSON table data and metadata from content.
    Returns dict with 'table_data' and 'summary' keys if found, None otherwise.
    """
    import re
    
    # First, try to parse entire content as JSON
    try:
        data = json.loads(content.strip())
        return extract_table_from_json(data)
    except:
        pass
    
    # Try to find JSON objects in the content
    # Look for JSON object boundaries
    json_start = content.find('{')
    if json_start != -1:
        # Try to find matching closing brace
        brace_count = 0
        json_end = -1
        for i in range(json_start, len(content)):
            if content[i] == '{':
                brace_count += 1
            elif content[i] == '}':
                brace_count -= 1
                if brace_count == 0:
                    json_end = i + 1
                    break
        
        if json_end > json_start:
            try:
                json_str = content[json_start:json_end]
                data = json.loads(json_str)
                result = extract_table_from_json(data)
                if result:
                    return result
            except:
                pass
    
    # Try to find JSON with "body" key containing JSON string
    body_match = re.search(r'"body"\s*:\s*"([^"]*(?:\\.[^"]*)*)"', content, re.DOTALL)
    if body_match:
        try:
            body_str = body_match.group(1)
            # Unescape JSON string
            body_str = body_str.replace('\\"', '"').replace('\\n', '\n').replace('\\t', '\t').replace('\\r', '\r')
            body_data = json.loads(body_str)
            result = extract_table_from_json(body_data)
            if result:
                return result
        except:
            pass
    
    # Try to find table_data pattern directly
    table_data_match = re.search(r'"table_data"\s*:\s*(\[[^\]]*(?:\[[^\]]*\][^\]]*)*\])', content, re.DOTALL)
    if table_data_match:
        try:
            table_str = table_data_match.group(1)
            table_data = json.loads(table_str)
            if isinstance(table_data, list):
                return {'table_data': table_data}
        except:
            pass
    
    return None

def extract_table_from_json(data):
    """
    Extract table_data and summary from a JSON structure.
    Returns dict with 'table_data' and 'summary' keys if found, None otherwise.
    """
    result = {}
    
    if isinstance(data, dict):
        if 'table_data' in data:
            result['table_data'] = data['table_data']
        if 'summary' in data:
            result['summary'] = data['summary']
        
        # Check nested body
        if 'body' in data and isinstance(data['body'], str):
            try:
                body_data = json.loads(data['body'])
                nested_result = extract_table_from_json(body_data)
                if nested_result:
                    result.update(nested_result)
            except:
                pass
        
        return result if result else None
    elif isinstance(data, list) and len(data) > 0:
        if isinstance(data[0], dict):
            return {'table_data': data}
    
    return None

def remove_json_from_content(content):
    """
    Remove JSON table data from content string.
    """
    import re
    
    # Remove JSON objects with table_data
    json_pattern = r'\{[^{}]*"table_data"\s*:\s*\[[^\]]*\][^{}]*\}'
    content = re.sub(json_pattern, '', content, flags=re.DOTALL)
    
    # Remove nested JSON body patterns
    nested_pattern = r'"body"\s*:\s*"[^"]*"'
    content = re.sub(nested_pattern, '', content, flags=re.DOTALL)
    
    return content.strip()

def create_word_table(doc, table_data):
    """
    Create a Word table from table_data (list of dictionaries).
    """
    if not table_data or len(table_data) == 0:
        return
    
    if not isinstance(table_data, list):
        return
    
    # Get all unique keys from all rows to determine columns
    all_keys = set()
    for row in table_data:
        if isinstance(row, dict):
            all_keys.update(row.keys())
    
    if not all_keys:
        return
    
    # Convert to sorted list for consistent ordering
    headers = sorted(all_keys)
    
    # Create table with headers + data rows
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header).replace('_', ' ').title()
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add data rows
    for row_data in table_data:
        if isinstance(row_data, dict):
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                value = str(row_data.get(header, ''))
                # Clean up value (remove markdown formatting)
                value = value.replace('**', '').strip()
                row_cells[i].text = value

def add_inline_formatting(paragraph, text):
    """
    Add inline formatting (bold, etc.) to paragraph text.
    Handles **bold** markdown-style formatting.
    """
    import re
    
    if not text:
        return
    
    # Split text by **bold** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            bold_text = part[2:-2]  # Remove ** markers
            run = paragraph.add_run(bold_text)
            run.bold = True
        else:
            # Regular text
            if part:
                paragraph.add_run(part)

if __name__ == "__main__":
    main()

